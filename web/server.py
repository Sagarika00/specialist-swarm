"""
Deal Desk Copilot — web backend.

A thin FastAPI layer over the specialist swarm that adds the things a business
stakeholder actually needs: a live cost/usage meter, a compliance scan of the
deliverable, a rendered proposal viewer, and a run history / audit trail.

Endpoints
    GET  /                       the console UI
    GET  /proposal/{sid}         rendered proposal viewer
    POST /api/run                start a run  {document, mode: "fast"|"full"}
    GET  /api/stream/{sid}       SSE: live swarm events + running token meter
    GET  /api/summary/{sid}      status, timings, usage, cost, specialist verdicts
    GET  /api/deliverable/{sid}  the client-facing deliverable, rendered to HTML
    GET  /api/files/{sid}        raw file list
    GET  /api/files/{sid}/{fid}  download one raw file
    GET  /api/scan/{sid}         compliance scan of the deliverable
    GET  /api/history            past runs

Run:
    cd web && python -m uvicorn server:app --reload --port 8000
"""

from __future__ import annotations

import json
import queue
import re
import threading
import time
from datetime import datetime, timezone
from pathlib import Path
from typing import Optional

import markdown as md
from anthropic import Anthropic
from dotenv import load_dotenv
from fastapi import FastAPI, File, HTTPException, UploadFile
from fastapi.responses import FileResponse, HTMLResponse, StreamingResponse
from pydantic import BaseModel

ROOT = Path(__file__).resolve().parent.parent
WEB = Path(__file__).resolve().parent
load_dotenv(ROOT / ".env")

BETA = "managed-agents-2026-04-01"
OUTPUT_DIR = ROOT / "outputs"
HISTORY_FILE = WEB / ".runs.json"

# Rough per-Mtoken USD rates for the LIVE estimate only. The authoritative
# number is the API's own `session.usage.list_cost`, shown alongside.
RATES = {
    "input": 3.0, "output": 15.0, "cache_write": 3.75, "cache_read": 0.30,
}

app = FastAPI(title="Deal Desk Copilot")
client = Anthropic(default_headers={"anthropic-beta": BETA})

# in-memory run registry: sid -> dict
_runs: dict[str, dict] = {}


# ────────────────────────────────────────────────────────────── helpers ──
def _ids() -> tuple[str, str]:
    cid, eid = ROOT / ".coordinator_id", ROOT / ".environment_id"
    if not cid.exists() or not eid.exists():
        raise HTTPException(500, "Swarm not built. Run setup_environment.py, "
                                 "create_specialists.py, upload_skills.py, create_coordinator.py.")
    return cid.read_text().strip(), eid.read_text().strip()


def _load_history() -> list[dict]:
    if HISTORY_FILE.exists():
        try:
            return json.loads(HISTORY_FILE.read_text())
        except json.JSONDecodeError:
            return []
    return []


def _save_history(rows: list[dict]) -> None:
    HISTORY_FILE.write_text(json.dumps(rows[-100:], indent=2))


def _history_upsert(sid: str, **fields) -> None:
    rows = _load_history()
    for r in rows:
        if r["sid"] == sid:
            for k, v in fields.items():
                if k == "created" and r.get("created"):
                    continue  # never move a run's original timestamp
                r[k] = v
            break
    else:
        rows.append({"sid": sid, **fields})
    _save_history(rows)


def _title_from(document: str) -> str:
    for line in document.splitlines():
        s = line.strip().lstrip("#").strip()
        if s:
            return s[:80]
    return "Untitled run"


def _estimate_cost(u: dict) -> float:
    return round(
        u.get("input", 0) * RATES["input"] / 1e6
        + u.get("output", 0) * RATES["output"] / 1e6
        + u.get("cache_write", 0) * RATES["cache_write"] / 1e6
        + u.get("cache_read", 0) * RATES["cache_read"] / 1e6,
        4,
    )


CHAT_MODEL = "claude-sonnet-5"


# ─────────────────────────────────────────────────────────────── models ──
class RunRequest(BaseModel):
    document: str
    mode: str = "fast"          # "fast" (markdown, ~5 min) | "full" (branded .docx)
    session_id: Optional[str] = None   # reuse a session created by /api/analyze


class AnalyzeRequest(BaseModel):
    document: str


class ChatRequest(BaseModel):
    message: str


def _new_run(document: str, mode: str = "fast", status: str = "analyzing") -> tuple[str, dict]:
    session = client.beta.sessions.create(
        agent=_ids()[0], environment_id=_ids()[1], title="Deal Desk Copilot",
    )
    run = {
        "document": document, "mode": mode, "t0": time.time(),
        "usage": {}, "specialists": {}, "status": status, "coordinator_text": "",
        "sse_log": [], "listeners": [], "lock": threading.Lock(),
        "chat": [], "brief": None,
    }
    run["sid"] = session.id
    _runs[session.id] = run
    (ROOT / ".last_session_id").write_text(session.id)
    return session.id, run


def _context_blocks(run: dict, include_deliverable: bool = True) -> str:
    """Everything the chat / brief models are allowed to reason from."""
    parts = [f"===== RFP DOCUMENT =====\n{run.get('document','')}"]
    for name, text in run.get("specialists", {}).items():
        parts.append(f"===== SPECIALIST VERDICT — {name} =====\n{text}")
    if include_deliverable and run.get("status") == "done":
        dtext = run.get("deliverable_cache")
        if dtext is None:
            try:
                _, dtext, _ = _deliverable_text(run["sid"])
            except Exception:
                dtext = ""
            run["deliverable_cache"] = dtext
        dtext = dtext or run.get("coordinator_text", "")
        if dtext:
            parts.append(f"===== CURRENT PROPOSAL DRAFT =====\n{dtext[:16000]}")
    return "\n\n".join(parts)


# ──────────────────────────────────────────────────────────────── pages ──
@app.get("/", response_class=HTMLResponse)
def index() -> str:
    return (WEB / "index.html").read_text()


@app.get("/proposal/{sid}", response_class=HTMLResponse)
def proposal_page(sid: str) -> str:
    return (WEB / "proposal.html").read_text()


# ─────────────────────────────────────────────────────── analyze / run ──
@app.post("/api/analyze")
def analyze(req: AnalyzeRequest) -> dict:
    """Register an RFP without running the swarm — enables chat + instant read."""
    _ids()
    if not req.document.strip():
        raise HTTPException(400, "document is empty")
    sid, _ = _new_run(req.document, status="analyzing")
    _history_upsert(sid, title=_title_from(req.document), mode="—", status="analyzing",
                    created=datetime.now(timezone.utc).isoformat(timespec="seconds"), cost=None)
    return {"session_id": sid}


@app.post("/api/run")
def run(req: RunRequest) -> dict:
    _ids()
    if not req.document.strip():
        raise HTTPException(400, "document is empty")
    mode = "full" if req.mode == "full" else "fast"

    run = _runs.get(req.session_id) if req.session_id else None
    if run:
        sid = req.session_id
        run.update(document=req.document, mode=mode, status="running",
                   t0=time.time(), sse_log=[], deliverable_cache=None)
    else:
        sid, run = _new_run(req.document, mode, status="running")

    _history_upsert(sid, title=_title_from(req.document), mode=mode, status="running",
                    created=datetime.now(timezone.utc).isoformat(timespec="seconds"))
    threading.Thread(target=_pump, args=(sid,), daemon=True).start()
    return {"session_id": sid, "mode": mode}


def _emit(run: dict, event: str, data: dict) -> None:
    """Record a display event and fan it out to every connected SSE client."""
    frame = (event, data)
    with run["lock"]:
        run["sse_log"].append(frame)
        for lq in list(run["listeners"]):
            lq.put(frame)


def _pump(sid: str) -> None:
    """Owns the run: consumes the raw event stream, updates state, emits frames.
    Runs to completion regardless of whether any browser is watching."""
    run = _runs[sid]
    u = run["usage"]
    try:
        with client.beta.sessions.events.stream(sid) as stream:
            client.beta.sessions.events.send(sid, events=[{
                "type": "user.message",
                "content": [{"type": "text",
                             "text": _user_message(run["document"], run["mode"])}],
            }])
            _emit(run, "open", {"session_id": sid, "mode": run["mode"]})
            for e in stream:
                t = e.type
                if t == "session.thread_created":
                    _emit(run, "thread_created", {"agent": getattr(e, "agent_name", "?")})
                elif t == "session.thread_status_running":
                    n = getattr(e, "agent_name", "?")
                    if n != "Deal Desk Senior Partner":
                        _emit(run, "thread_running", {"agent": n})
                elif t == "agent.thread_message_sent":
                    _emit(run, "delegate", {"to": getattr(e, "to_agent_name", None) or "specialist"})
                elif t == "agent.thread_message_received":
                    n = getattr(e, "from_agent_name", "?")
                    txt = "".join(getattr(b, "text", "") for b in getattr(e, "content", [])
                                  if getattr(b, "type", None) == "text")
                    run["specialists"][n] = txt
                    _emit(run, "reply", {"from": n, "text": txt})
                elif t == "agent.tool_use":
                    u["tools"] = u.get("tools", 0) + 1
                    _emit(run, "tool", {"name": getattr(e, "name", "?"), "tools": u["tools"]})
                elif t == "agent.message":
                    txt = "".join(getattr(b, "text", "") for b in getattr(e, "content", [])
                                  if getattr(b, "type", None) == "text")
                    if txt:
                        run["coordinator_text"] += txt + "\n"
                        _emit(run, "message", {"text": txt})
                elif t == "span.model_request_end":
                    mu = getattr(e, "model_usage", None)
                    if mu:
                        d = mu.model_dump() if hasattr(mu, "model_dump") else dict(mu)
                        u["input"] = u.get("input", 0) + d.get("input_tokens", 0)
                        u["output"] = u.get("output", 0) + d.get("output_tokens", 0)
                        u["cache_write"] = u.get("cache_write", 0) + d.get("cache_creation_input_tokens", 0)
                        u["cache_read"] = u.get("cache_read", 0) + d.get("cache_read_input_tokens", 0)
                        u["requests"] = u.get("requests", 0) + 1
                        _emit(run, "usage", {**u, "est_cost": _estimate_cost(u),
                                             "elapsed": round(time.time() - run["t0"])})
                elif t == "session.usage":
                    us = getattr(e, "usage", None)
                    if us:
                        d = us.model_dump() if hasattr(us, "model_dump") else dict(us)
                        lc = d.get("list_cost") or {}
                        run["api_cost"] = lc.get("amount")
                        run["api_currency"] = lc.get("currency", "USD")
                        run["active_seconds"] = d.get("active_seconds")
                elif t == "session.status_idle":
                    break
        run["status"] = "done"
        _history_upsert(sid, status="done", cost=run.get("api_cost") or _estimate_cost(u))
        _emit(run, "done", {"est_cost": _estimate_cost(u), "api_cost": run.get("api_cost"),
                            "api_currency": run.get("api_currency", "USD")})
    except Exception as exc:
        run["status"] = "error"
        _history_upsert(sid, status="error")
        _emit(run, "error", {"message": str(exc)})
    finally:
        _emit(run, "close", {})
        run["done_streaming"] = True


def _user_message(document: str, mode: str) -> str:
    fmt = ("Produce the final deliverable as a **markdown** file named "
           "proposal-response.md. Do NOT use the docx skill."
           if mode == "fast" else
           "Produce the final deliverable as a branded **Word document** via the docx skill.")
    return (
        "An RFP has landed. Run the standard Deal Desk process, MOVE FAST:\n"
        "1. Read the RFP below — do not explore the filesystem.\n"
        "2. Delegate to ALL FOUR specialists in ONE turn, parallel.\n"
        "3. Synthesise their replies.\n"
        f"4. {fmt} Also write a second `INTERNAL_` file with the internal-only brief.\n\n"
        f"===== RFP DOCUMENT =====\n{document}"
    )


def _sse(event: str, data: dict) -> str:
    return f"event: {event}\ndata: {json.dumps(data, default=str)}\n\n"


def _event_stream(sid: str):
    run = _runs.get(sid)
    if not run:
        yield _sse("error", {"message": "unknown session — start a new run"}); return

    lq: queue.Queue = queue.Queue()
    with run["lock"]:
        backlog = list(run["sse_log"])
        if not run.get("done_streaming"):
            run["listeners"].append(lq)
    # replay everything so far (reconnect / late-open safe)
    for event, data in backlog:
        yield _sse(event, data)
    if run.get("done_streaming"):
        return
    try:
        while True:
            event, data = lq.get()
            yield _sse(event, data)
            if event in ("close",):
                break
    finally:
        with run["lock"]:
            if lq in run["listeners"]:
                run["listeners"].remove(lq)


@app.get("/api/stream/{sid}")
def stream(sid: str):
    return StreamingResponse(_event_stream(sid), media_type="text/event-stream",
                             headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"})


# ──────────────────────────────────────────────────────── summary/files ──
@app.get("/api/summary/{sid}")
def summary(sid: str) -> dict:
    run = _runs.get(sid, {})
    u = run.get("usage", {})
    return {
        "status": run.get("status", "unknown"),
        "mode": run.get("mode"),
        "elapsed": round(time.time() - run["t0"]) if run.get("t0") else None,
        "active_seconds": run.get("active_seconds"),
        "usage": u,
        "tool_calls": u.get("tools", 0),
        "est_cost": _estimate_cost(u),
        "api_cost": run.get("api_cost"),
        "api_currency": run.get("api_currency", "USD"),
        "rates": RATES,
        "specialists": [{"name": k, "text": v} for k, v in run.get("specialists", {}).items()],
    }


@app.get("/api/files/{sid}")
def files(sid: str) -> dict:
    listing = client.beta.files.list(scope_id=sid, betas=[BETA])
    return {"files": [{"id": f.id, "filename": f.filename,
                       "internal": "INTERNAL" in f.filename.upper()} for f in listing.data]}


@app.get("/api/files/{sid}/{fid}")
def download(sid: str, fid: str):
    listing = client.beta.files.list(scope_id=sid, betas=[BETA])
    match = next((f for f in listing.data if f.id == fid), None)
    if not match:
        raise HTTPException(404, "file not found")
    OUTPUT_DIR.mkdir(exist_ok=True)
    out = OUTPUT_DIR / match.filename
    client.beta.files.download(fid).write_to_file(str(out))
    return FileResponse(out, filename=match.filename)


def _deliverable_text(sid: str) -> tuple[str, str, str]:
    """Return (filename, raw_text, markdown_or_plain). Prefers client-facing file."""
    listing = list(client.beta.files.list(scope_id=sid, betas=[BETA]).data)
    if not listing:
        return ("", "", "")
    external = [f for f in listing if "INTERNAL" not in f.filename.upper()]
    pool = external or listing
    pick = next((f for f in pool if f.filename.lower().endswith(".md")), pool[0])
    OUTPUT_DIR.mkdir(exist_ok=True)
    out = OUTPUT_DIR / pick.filename
    client.beta.files.download(pick.id).write_to_file(str(out))
    if pick.filename.lower().endswith(".docx"):
        from docx import Document
        doc = Document(str(out))
        text = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
    else:
        text = out.read_text(errors="replace")
    return (pick.filename, text, text)


@app.get("/api/deliverable/{sid}")
def deliverable(sid: str) -> dict:
    run = _runs.get(sid, {})
    filename, text, body = _deliverable_text(sid)
    if not text and run.get("coordinator_text"):
        filename, text, body = ("coordinator-response.md", run["coordinator_text"],
                                run["coordinator_text"])
    html = md.markdown(body, extensions=["tables", "fenced_code", "toc", "sane_lists"])
    all_files = files(sid)["files"]
    return {
        "filename": filename, "html": html, "markdown": text,
        "title": run.get("document", "").splitlines()[0].lstrip("#").strip()
                 if run.get("document") else filename,
        "files": all_files,
    }


# ─────────────────────────────────────────────────────── compliance scan ──
SCAN_RULES = [
    ("critical", "Private key", re.compile(r"-----BEGIN [A-Z ]*PRIVATE KEY-----")),
    ("critical", "Anthropic API key", re.compile(r"sk-ant-[A-Za-z0-9_\-]{20,}")),
    ("critical", "AWS access key", re.compile(r"AKIA[0-9A-Z]{16}")),
    ("critical", "GitHub token", re.compile(r"gh[pousr]_[A-Za-z0-9]{30,}")),
    ("critical", "Generic secret assignment",
     re.compile(r"(?i)\b(api[_-]?key|secret|passwd|password|access[_-]?token|bearer)\b\s*[:=]\s*['\"]?[A-Za-z0-9/\+_\-]{8,}")),
    ("high", "Internal-only content in client deliverable",
     re.compile(r"(?i)\b(walk-?away|do not send|internal only|list price|gross margin|price floor|red-?line|MFN strategy)\b")),
    ("medium", "Social Security Number", re.compile(r"\b\d{3}-\d{2}-\d{4}\b")),
    ("medium", "Payment card number", re.compile(r"\b(?:\d[ -]?){13,16}\b")),
    ("low", "Email address", re.compile(r"\b[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}\b")),
    ("low", "Phone number",
     re.compile(r"(?<!\d)(?:\+\d[\d ().\-]{7,}\d|\(\d{3}\)[\d .\-]{6,}\d|\d{3}[ .\-]\d{3}[ .\-]\d{4})(?!\d)")),
]
_SEV_RANK = {"critical": 4, "high": 3, "medium": 2, "low": 1}
_ISO_DATE = re.compile(r"^\d{4}-\d{2}-\d{2}$")


def _luhn_ok(s: str) -> bool:
    digs = [int(c) for c in s if c.isdigit()]
    if not 13 <= len(digs) <= 19:
        return False
    chk = 0
    for i, d in enumerate(reversed(digs)):
        if i % 2:
            d *= 2
            if d > 9:
                d -= 9
        chk += d
    return chk % 10 == 0


@app.get("/api/scan/{sid}")
def scan(sid: str) -> dict:
    _, text, _ = _deliverable_text(sid)
    if not text:
        text = _runs.get(sid, {}).get("coordinator_text", "")
    lines = text.splitlines()
    findings: list[dict] = []
    for sev, label, rx in SCAN_RULES:
        for m in rx.finditer(text):
            raw = m.group(0)
            if _ISO_DATE.match(raw.strip()):
                continue
            if label == "Payment card number" and not _luhn_ok(raw):
                continue
            ln = text.count("\n", 0, m.start()) + 1
            snippet = raw if len(raw) <= 8 else raw[:3] + "…" + raw[-3:]
            if sev == "high":
                snippet = raw  # internal-content matches are safe to show
            findings.append({"severity": sev, "type": label, "line": ln,
                             "match": snippet,
                             "context": lines[ln - 1].strip()[:160] if ln <= len(lines) else ""})
    findings.sort(key=lambda f: (-_SEV_RANK[f["severity"]], f["line"]))
    worst = max((_SEV_RANK[f["severity"]] for f in findings), default=0)
    verdict = ("fail" if worst >= 3 else "warn" if worst == 2 else
               "review" if worst == 1 else "pass")
    counts = {s: sum(1 for f in findings if f["severity"] == s)
              for s in ("critical", "high", "medium", "low")}
    return {"verdict": verdict, "counts": counts, "findings": findings[:100],
            "scanned_chars": len(text)}


@app.get("/api/history")
def history() -> dict:
    return {"runs": list(reversed(_load_history()))}


@app.get("/api/sample")
def sample() -> dict:
    p = ROOT / "synthetic-data" / "rfp-acme-corp.md"
    return {"document": p.read_text() if p.exists() else ""}


# ─────────────────────────────────────────── upload & parse a document ──
@app.post("/api/upload")
async def upload(file: UploadFile = File(...)) -> dict:
    raw = await file.read()
    name = (file.filename or "document").lower()
    if len(raw) > 12_000_000:
        raise HTTPException(413, "file too large (12 MB max)")
    try:
        if name.endswith(".pdf"):
            import io
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(raw))
            text = "\n\n".join(p.extract_text() or "" for p in reader.pages)
        elif name.endswith(".docx"):
            import io
            from docx import Document
            doc = Document(io.BytesIO(raw))
            text = "\n".join(p.text for p in doc.paragraphs)
        elif name.endswith((".txt", ".md", ".markdown", ".rtf")):
            text = raw.decode("utf-8", errors="replace")
        else:
            raise HTTPException(415, f"unsupported type: {name.rsplit('.', 1)[-1]}. "
                                     "Use PDF, DOCX, TXT or MD, or paste the text.")
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(422, f"could not parse {file.filename}: {exc}")
    text = text.strip()
    if len(text) < 40:
        raise HTTPException(422, "extracted almost no text — is this a scanned PDF? Paste the text instead.")
    return {"document": text, "filename": file.filename, "chars": len(text)}


# ─────────────────────────────────── proof points: mandatory reqs ↔ track record ──
def _load_json(name: str) -> dict:
    p = ROOT / "synthetic-data" / name
    return json.loads(p.read_text()) if p.exists() else {}


PROOF_SYSTEM = (
    "You are a proposal manager building the proof-point matrix for an RFP response.\n"
    "You get: (1) an RFP, (2) our firm's delivered-engagements track record as JSON.\n"
    "Extract every MANDATORY / non-negotiable requirement from the RFP (things the vendor "
    "MUST do — capabilities, integrations, residency, scale, SLAs, commercial structure). "
    "For each, find the engagement(s) in the track record that prove we've done it before.\n"
    "Return STRICT JSON only:\n"
    '{"requirements":[{"requirement":str,"category":str,'
    '"status":"proven"|"partial"|"gap",'
    '"matches":[{"engagement_id":str,"project":str,"client":str,"year":int,'
    '"strength":"strong"|"moderate","evidence":str,"reference_available":bool}],'
    '"note":str}],'
    '"summary":{"proven":int,"partial":int,"gap":int}}\n'
    "status=proven if a strong match exists; partial if only adjacent experience; gap if nothing. "
    "Never invent engagements — only use ones present in the track record JSON."
)


@app.post("/api/proofpoints/{sid}")
def proofpoints(sid: str) -> dict:
    run = _runs.get(sid)
    if not run:
        raise HTTPException(404, "unknown session")
    if run.get("proofpoints"):
        return run["proofpoints"]
    track = _load_json("past-engagements.json")
    resp = client.messages.create(
        model=CHAT_MODEL, max_tokens=8000, system=PROOF_SYSTEM,
        messages=[{"role": "user", "content":
                   f"===== RFP =====\n{run['document'][:24000]}\n\n"
                   f"===== OUR TRACK RECORD (JSON) =====\n{json.dumps(track)}"}],
    )
    data = _extract_json(resp)
    reqs = data.get("requirements", [])
    data["summary"] = {
        s: sum(1 for r in reqs if r.get("status") == s) for s in ("proven", "partial", "gap")
    }
    run["proofpoints"] = data
    return data


# ───────────────────────────────────────── win + approval scorecard ──
FACTOR_WEIGHTS = {
    "solution_fit": 0.25,
    "competitive_position": 0.20,
    "relationship_and_proof": 0.15,
    "commercial_alignment": 0.20,
    "risk_and_compliance": 0.10,
    "deadline_feasibility": 0.10,
}

SCORE_SYSTEM = (
    "You are a bid-desk analyst scoring an opportunity. You get the RFP, our pricing "
    "playbook rules, our past wins, and our delivered-engagements track record.\n"
    "Score each factor 0-100 (100 = strongly in our favour) with a one-line rationale. "
    "Also estimate the commercial concessions the RFP would force and which internal "
    "approvers that triggers.\n"
    "Return STRICT JSON only:\n"
    '{"factors":{'
    '"solution_fit":{"score":int,"why":str},'
    '"competitive_position":{"score":int,"why":str},'
    '"relationship_and_proof":{"score":int,"why":str},'
    '"commercial_alignment":{"score":int,"why":str},'
    '"risk_and_compliance":{"score":int,"why":str},'
    '"deadline_feasibility":{"score":int,"why":str}},'
    '"approval":{"probability":int,"approvers_required":[str],'
    '"drivers":[{"item":str,"effect":"+"|"-","note":str}]},'
    '"segment":{"name":str,"historical_win_rate":int,"basis":str},'
    '"headline":str}\n'
    "approval.probability = chance ALL required internal approvers sign off given the "
    "concessions needed (deeper discount than band, non-standard MSA, sub-band margin, "
    "uninsurable liability). Be realistic — stacked blockers lower it."
)


@app.post("/api/scorecard/{sid}")
def scorecard(sid: str) -> dict:
    run = _runs.get(sid)
    if not run:
        raise HTTPException(404, "unknown session")
    if run.get("scorecard"):
        return run["scorecard"]

    playbook = (ROOT / "skills" / "pricing-playbook" / "SKILL.md")
    ctx = [f"===== RFP =====\n{run['document'][:20000]}"]
    if playbook.exists():
        ctx.append(f"===== PRICING PLAYBOOK =====\n{playbook.read_text()[:6000]}")
    ctx.append(f"===== PAST WINS =====\n{json.dumps(_load_json('past-wins.json'))}")
    ctx.append(f"===== TRACK RECORD =====\n{json.dumps(_load_json('past-engagements.json'))}")
    for name, text in run.get("specialists", {}).items():
        ctx.append(f"===== SPECIALIST VERDICT — {name} =====\n{text}")

    resp = client.messages.create(
        model=CHAT_MODEL, max_tokens=4000, system=SCORE_SYSTEM,
        messages=[{"role": "user", "content": "\n\n".join(ctx)}],
    )
    data = _extract_json(resp)

    factors = data.get("factors", {})
    win = 0.0
    breakdown = []
    for key, weight in FACTOR_WEIGHTS.items():
        s = float(factors.get(key, {}).get("score", 50))
        contribution = s * weight
        win += contribution
        breakdown.append({
            "name": key.replace("_", " "), "score": round(s), "weight": weight,
            "contribution": round(contribution, 1),
            "why": factors.get(key, {}).get("why", ""),
        })
    data["win_probability"] = round(win)
    data["breakdown"] = sorted(breakdown, key=lambda b: -b["contribution"])
    data["weights"] = FACTOR_WEIGHTS
    run["scorecard"] = data
    return data


def _extract_json(resp) -> dict:
    raw = "".join(b.text for b in resp.content if getattr(b, "type", None) == "text")
    raw = raw.replace("```json", "").replace("```", "").strip()
    m = re.search(r"\{.*\}", raw, re.S)
    try:
        return json.loads(m.group(0) if m else raw)
    except Exception:
        return {"error": "model did not return valid JSON", "raw": raw[:800]}


# ─────────────────────────────────────────────────── chat over the RFP ──
CHAT_SYSTEM = (
    "You are the Deal Desk Analyst — an assistant to a partner working an inbound RFP.\n"
    "Answer ONLY from the materials provided below (the RFP, any specialist verdicts, "
    "and the current proposal draft). Be specific: cite numbers, clauses, and section "
    "names. Lead with the answer. Flag risk plainly. If the materials don't contain the "
    "answer, say so and say which specialist would know.\n"
    "Keep answers under 250 words unless asked to draft something."
)


@app.post("/api/chat/{sid}")
def chat(sid: str, req: ChatRequest):
    run = _runs.get(sid)
    if not run:
        raise HTTPException(404, "unknown session — analyse or run an RFP first")
    msg = req.message.strip()
    if not msg:
        raise HTTPException(400, "empty message")

    history = run.setdefault("chat", [])
    history.append({"role": "user", "content": msg})
    context = _context_blocks(run)
    messages = (
        [{"role": "user", "content": f"MATERIALS:\n\n{context}"},
         {"role": "assistant", "content": "Understood. I have the RFP"
          + (" and the deal team's analysis" if run.get("specialists") else "")
          + ". What do you need?"}]
        + history[-12:]
    )

    def gen():
        acc = []
        try:
            with client.messages.stream(model=CHAT_MODEL, max_tokens=1200,
                                        system=CHAT_SYSTEM, messages=messages) as s:
                for tok in s.text_stream:
                    acc.append(tok)
                    yield tok
        except Exception as exc:
            yield f"\n\n[chat error: {exc}]"
        finally:
            history.append({"role": "assistant", "content": "".join(acc)})

    return StreamingResponse(gen(), media_type="text/plain")


@app.get("/api/chat/{sid}")
def chat_history(sid: str) -> dict:
    return {"messages": _runs.get(sid, {}).get("chat", [])}


# ───────────────────────────────────────── instant read (pre-swarm) ──
BRIEF_SYSTEM = (
    "You are a bid manager doing a 30-second triage of an inbound RFP. Return STRICT "
    "JSON only, no prose, matching exactly this shape:\n"
    '{"customer": str, "deal_size_estimate": str, "response_due": str, '
    '"mandatory_requirements": int, "evaluation_criteria": [str], '
    '"red_flags": [{"item": str, "why": str}], '
    '"bid_lean": "BID" | "BID WITH CONDITIONS" | "NO-BID", "rationale": str}\n'
    "red_flags: the 3-5 most dangerous asks (uncapped liability, MFN, price-only eval, "
    "unwinnable incumbent, etc). Be blunt in rationale."
)


@app.post("/api/brief/{sid}")
def brief(sid: str) -> dict:
    run = _runs.get(sid)
    if not run:
        raise HTTPException(404, "unknown session")
    if run.get("brief"):
        return run["brief"]
    resp = client.messages.create(
        model=CHAT_MODEL, max_tokens=4000, system=BRIEF_SYSTEM,
        messages=[{"role": "user", "content": run["document"][:24000]}],
    )
    raw = "".join(b.text for b in resp.content if getattr(b, "type", None) == "text")
    raw = raw.replace("```json", "").replace("```", "").strip()
    m = re.search(r"\{.*\}", raw, re.S)
    try:
        data = json.loads(m.group(0) if m else raw)
    except Exception:
        data = {"customer": "?", "rationale": raw[:500], "bid_lean": "BID WITH CONDITIONS",
                "red_flags": [], "evaluation_criteria": [], "mandatory_requirements": 0,
                "deal_size_estimate": "?", "response_due": "?"}
    run["brief"] = data
    return data
